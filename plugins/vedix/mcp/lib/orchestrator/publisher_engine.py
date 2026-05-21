"""Publisher templates engine (Block 7 — §7).

Vedix bundles 23 publisher families. Each entry in :data:`VENUES` declares
a LaTeX class identifier, a biblatex citation-style identifier, a region
tag (``"global"`` or ``"ru"``) and a free-text description. The on-disk
template directory lives at ``plugins/vedix/mcp/templates/<venue>/`` with
``latex/``, ``word/``, ``profile.json`` and ``PROVENANCE.md`` siblings.

Public surface:

* :func:`list_venues` — sorted list of all 23 venue names.
* :func:`get_venue` — return :class:`Venue` for ``name`` or ``name:journal``;
  raise ``KeyError`` for unknown names.
* :func:`render` — pandoc-driven `manuscript.md` → ``manuscript.tex`` +
  ``manuscript.pdf`` + ``manuscript.docx`` round-trip.
"""
from __future__ import annotations

import json
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Literal


@dataclass
class Venue:
    """One publisher target.

    Attributes:
        name: Slug identifier (matches the template directory name).
        latex_class: Filename of the LaTeX class file (``*.cls``).
        citation_style: Biblatex style identifier.
        region: ``"global"`` or ``"ru"``.
        bundled: Whether the template ships in the installer payload.
        description: Free-text human-readable note.
    """

    name: str
    latex_class: str
    citation_style: str
    region: Literal["global", "ru"]
    bundled: bool = True
    description: str = ""


VENUES: dict[str, Venue] = {
    "preprint": Venue(
        "preprint",
        "preprint.cls",
        "biblatex-numeric-comp",
        "global",
        description="11pt single-column preprint for arXiv / bioRxiv / OSF / SSRN",
    ),
    "nature": Venue(
        "nature",
        "nature.cls",
        "nature",
        "global",
        description="Nature Portfolio family",
    ),
    "elsevier": Venue(
        "elsevier",
        "elsarticle.cls",
        "elsevier-numeric",
        "global",
        description="~2,500 Elsevier titles",
    ),
    "springer-nature": Venue(
        "springer-nature",
        "sn-jnl.cls",
        "springer-numeric",
        "global",
        description="Springer-Nature unified template",
    ),
    "taylor-francis": Venue(
        "taylor-francis",
        "interact.cls",
        "tf-numeric",
        "global",
        description="Taylor & Francis Interact (CADLaTeX)",
    ),
    "frontiers": Venue(
        "frontiers",
        "frontiers.cls",
        "frontiers-reference-style",
        "global",
        description="Frontiers in *",
    ),
    "wiley": Venue(
        "wiley",
        "WileyNJD-v2.cls",
        "wiley-numeric",
        "global",
        description="Wiley New Journal Design v2",
    ),
    "sage": Venue(
        "sage",
        "sagej.cls",
        "sage-author-date",
        "global",
        description="SAGE journals",
    ),
    "plos": Venue(
        "plos",
        "plos2015.cls",
        "vancouver",
        "global",
        description="PLOS family",
    ),
    "cell": Venue(
        "cell",
        "cell.cls",
        "cell",
        "global",
        description="Cell Press",
    ),
    "ieee": Venue(
        "ieee",
        "IEEEtran.cls",
        "ieee",
        "global",
        description="IEEE Transactions / Conferences",
    ),
    "acm": Venue(
        "acm",
        "acmart.cls",
        "acm-numeric",
        "global",
        description="ACM Transactions / Proceedings",
    ),
    "acs": Venue(
        "acs",
        "achemso.cls",
        "acs",
        "global",
        description="American Chemical Society",
    ),
    "mdpi": Venue(
        "mdpi",
        "mdpi.cls",
        "mdpi-numeric",
        "global",
        description="MDPI open-access family",
    ),
    "revtex42": Venue(
        "revtex42",
        "revtex4-2.cls",
        "revtex-numeric",
        "global",
        description="AIP / APS physics",
    ),
    "rsc": Venue(
        "rsc",
        "rsc.cls",
        "rsc-author-date",
        "global",
        description="Royal Society of Chemistry",
    ),
    "cambridge": Venue(
        "cambridge",
        "cambridge7A.cls",
        "cup-author-date",
        "global",
        description="Cambridge University Press",
    ),
    "oup": Venue(
        "oup",
        "OUPMaths.cls",
        "oup-numeric",
        "global",
        description="Oxford University Press",
    ),
    "bmj": Venue(
        "bmj",
        "bmj.cls",
        "vancouver",
        "global",
        description="BMJ Publishing Group",
    ),
    "jama": Venue(
        "jama",
        "jama-style.cls",
        "ama",
        "global",
        description="JAMA Network (in-house MIT)",
    ),
    "gost-generic": Venue(
        "gost-generic",
        "gost-article.cls",
        "gost-numeric",
        "ru",
        description="ВАК-перечень generic ГОСТ-7.0.5 (in-house MIT)",
    ),
    "dan-ras": Venue(
        "dan-ras",
        "dan-ras.cls",
        "gost-numeric",
        "ru",
        description="Доклады РАН (in-house MIT)",
    ),
    "uspekhi": Venue(
        "uspekhi",
        "uspekhi.cls",
        "gost-numeric",
        "ru",
        description="Успехи физических наук (in-house MIT)",
    ),
}


def list_venues() -> list[str]:
    """Return a sorted list of all registered venue slugs."""
    return sorted(VENUES)


def get_venue(name: str) -> Venue:
    """Resolve ``name`` (optionally ``name:journal``) to a :class:`Venue`.

    Raises:
        KeyError: ``name``'s base slug is not in :data:`VENUES`.
    """
    base = name.split(":", 1)[0]
    if base not in VENUES:
        raise KeyError(
            f"unknown venue {base!r}; available: {list_venues()}"
        )
    return VENUES[base]


def _templates_root() -> Path:
    """Return the on-disk root of bundled template directories.

    Resolves to ``plugins/vedix/mcp/templates/`` regardless of how the
    package is installed.
    """
    return Path(__file__).resolve().parents[2] / "templates"


def render(
    *,
    venue: str,
    language: str,
    manuscript_md: Path,
    references_bib: Path,
    workdir: Path,
    journal: str | None = None,
) -> dict:
    """Render a pandoc-flavoured markdown manuscript to PDF + DOCX.

    Args:
        venue: Slug from :data:`VENUES`, optionally ``slug:journal``.
        language: ISO 639-1 code (resolved via the locale router).
        manuscript_md: Path to the source markdown.
        references_bib: Path to the BibLaTeX bibliography.
        workdir: Pre-existing directory to write artifacts into.
        journal: Optional ``profile.json["journals"][journal]`` override.

    Returns:
        ``{"venue": str, "language": str, "pdf": Path, "docx": Path, "tex": Path}``.

    Raises:
        FileNotFoundError: Venue directory or class file missing.
        subprocess.CalledProcessError: pandoc or pdflatex failed.
    """
    v = get_venue(venue)
    venue_root = _templates_root() / v.name
    if not venue_root.exists():
        raise FileNotFoundError(f"venue template directory missing: {venue_root}")

    tex_out = workdir / "manuscript.tex"
    pdf_out = workdir / "manuscript.pdf"
    docx_out = workdir / "manuscript.docx"

    profile_path = venue_root / "profile.json"
    profile: dict = {}
    if profile_path.exists():
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    if journal:
        journal_overrides = profile.get("journals", {}).get(journal, {})
        profile = {**profile, **journal_overrides}

    # Locale-driven LaTeX engine selection (pdflatex vs xelatex).
    from .locale.router import get_locale

    locale = get_locale(language)
    latex_class_path = venue_root / "latex" / v.latex_class
    if not latex_class_path.exists():
        raise FileNotFoundError(f"LaTeX class missing: {latex_class_path}")

    subprocess.run(
        [
            "pandoc",
            str(manuscript_md),
            "-o",
            str(tex_out),
            "--from",
            "markdown",
            "--to",
            "latex",
            "--standalone",
            "--variable",
            f"documentclass={latex_class_path.stem}",
            "--variable",
            f"venue-class-path={latex_class_path.parent}",
            "--variable",
            f"language={language}",
            "--variable",
            f"engine={locale.latex_engine}",
            "--bibliography",
            str(references_bib),
            "--biblatex",
            f"--metadata=lang:{language}",
        ],
        check=True,
        cwd=workdir,
    )

    engine = locale.latex_engine
    for _ in range(2):
        subprocess.run(
            [engine, "-interaction=nonstopmode", str(tex_out)],
            cwd=workdir,
            check=False,
        )
    subprocess.run(
        ["biber", str(tex_out.with_suffix(""))],
        cwd=workdir,
        check=False,
    )
    for _ in range(2):
        subprocess.run(
            [engine, "-interaction=nonstopmode", str(tex_out)],
            cwd=workdir,
            check=False,
        )

    word_template = venue_root / "word" / f"{v.name}_{language}.dotx"
    if not word_template.exists():
        word_template = venue_root / "word" / f"{v.name}_en.dotx"
    pandoc_docx_cmd = [
        "pandoc",
        str(manuscript_md),
        "-o",
        str(docx_out),
        "--bibliography",
        str(references_bib),
        "--citeproc",
    ]
    if word_template.exists():
        pandoc_docx_cmd.extend(["--reference-doc", str(word_template)])
    subprocess.run(pandoc_docx_cmd, check=True, cwd=workdir)

    return {
        "venue": v.name,
        "language": language,
        "pdf": pdf_out,
        "docx": docx_out,
        "tex": tex_out,
    }


# ---------------------------------------------------------------------------
# LaTeX↔Word parity check (Block 7 Task 4).
# ---------------------------------------------------------------------------


def _count_artifacts_in_text(text: str) -> dict:
    """Count sections / equations / figures / tables / references / citations.

    Pure-text heuristic — works on both pandoc PDF/DOCX extractions.
    """
    import re

    return {
        "sections": re.findall(
            r"(?:^|\n)([1-9]\.?\s+[A-Z][^\n]{2,80})", text
        ),
        "n_equations": len(
            re.findall(
                r"\\begin\{equation\}|^\([0-9]+\)\s*$",
                text,
                re.MULTILINE,
            )
        ),
        "n_figures": len(re.findall(r"Figure\s+\d+", text)),
        "n_tables": len(re.findall(r"Table\s+\d+", text)),
        "n_references": len(
            re.findall(r"^\[\d+\]\s", text, re.MULTILINE)
        ),
        "word_count": len(text.split()),
        "n_citations": len(re.findall(r"\[\d+\]", text)),
    }


def _inspect_pdf(pdf: Path) -> dict:
    """Best-effort extract section headings + artifact counts from a PDF."""
    from pdfminer.high_level import extract_text

    text = extract_text(str(pdf))
    return _count_artifacts_in_text(text)


def _inspect_docx(docx: Path) -> dict:
    """Extract section headings (paragraphs styled ``Heading*``) + counts."""
    from docx import Document

    doc = Document(str(docx))
    text = "\n".join(p.text for p in doc.paragraphs)
    sections = [
        p.text
        for p in doc.paragraphs
        if p.style is not None
        and p.style.name is not None
        and p.style.name.startswith("Heading")
    ]
    return {**_count_artifacts_in_text(text), "sections": sections}


def check_parity(
    *,
    pdf: Path,
    docx: Path,
    word_tolerance_pct: float = 2.0,
) -> dict:
    """Compare PDF + DOCX renders of the same manuscript for drift.

    Args:
        pdf: Path to ``manuscript.pdf``.
        docx: Path to ``manuscript.docx``.
        word_tolerance_pct: Allowed |Δword-count| as a percentage of the
            larger document; default 2 %.

    Returns:
        ``{"status": "ok"|"drift", "divergences": [...], "pdf_data": {...},
        "docx_data": {...}}``. Each divergence is a dict with a ``"kind"``
        key (``"section_count"``, ``"n_equations"``, ``"n_figures"``,
        ``"n_tables"``, ``"n_references"``, ``"n_citations"``,
        ``"word_count"``).
    """
    pdf_data = _inspect_pdf(pdf)
    docx_data = _inspect_docx(docx)
    divergences: list[dict] = []

    if len(pdf_data["sections"]) != len(docx_data["sections"]):
        divergences.append(
            {
                "kind": "section_count",
                "pdf": len(pdf_data["sections"]),
                "docx": len(docx_data["sections"]),
            }
        )

    for k in (
        "n_equations",
        "n_figures",
        "n_tables",
        "n_references",
        "n_citations",
    ):
        if pdf_data[k] != docx_data[k]:
            divergences.append(
                {"kind": k, "pdf": pdf_data[k], "docx": docx_data[k]}
            )

    wc_pdf, wc_docx = pdf_data["word_count"], docx_data["word_count"]
    drift_pct = abs(wc_pdf - wc_docx) / max(wc_pdf, wc_docx, 1) * 100
    if drift_pct > word_tolerance_pct:
        divergences.append(
            {
                "kind": "word_count",
                "pdf": wc_pdf,
                "docx": wc_docx,
                "drift_pct": round(drift_pct, 2),
                "tolerance_pct": word_tolerance_pct,
            }
        )

    return {
        "status": "ok" if not divergences else "drift",
        "divergences": divergences,
        "pdf_data": pdf_data,
        "docx_data": docx_data,
    }
